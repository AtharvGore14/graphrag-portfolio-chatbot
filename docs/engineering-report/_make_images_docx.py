"""Generate Word attachment: Images Needed for Documentation."""
from docx import Document
from docx.shared import Pt

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

doc.add_heading("Images needed for the documentation", 0)
p = doc.add_paragraph()
run = p.add_run("Portfolio GraphRAG Chatbot — Industry Project")
run.italic = True

doc.add_paragraph("Upload to: Overleaf → images/ folder")
doc.add_paragraph(
    "Project: https://www.overleaf.com/1656157384wxcgqvgpkpdx#c9a2af"
)
doc.add_paragraph(
    "Use these exact filenames (PNG preferred). "
    "Please avoid showing real API keys or passwords."
)
doc.add_paragraph(
    "After chapter consolidation the report needs 22 figures total "
    "(10 screenshots + 12 diagrams). Anything not listed below was cut — "
    "do not draw or capture it."
)

doc.add_heading("Already done (skip)", level=1)
t = doc.add_table(rows=2, cols=2)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Filename"
t.rows[0].cells[1].text = "Notes"
t.rows[1].cells[0].text = "01_college_logo.png"
t.rows[1].cells[1].text = "Already in latex/images/"

doc.add_heading("A) Screenshots — after running the project (10 required)", level=1)
shots = [
    ("1", "S01_compose_healthy.png", "Compose stack healthy",
     "docker compose up --build then Desktop / docker compose ps", "Ch 4"),
    ("2", "S02_health_json.png", "Health JSON",
     "Open http://localhost:8000/health", "Ch 4"),
    ("3", "S03_swagger_docs.png", "Swagger UI",
     "Open http://localhost:8000/docs", "Ch 6"),
    ("4", "S05_ui_grounded_answer.png", "Grounded answer + citations",
     "user_id=demo, ask e.g. ADANIPOWER holdings", "Ch 7"),
    ("5", "S06_ui_refuse.png", "Refuse / clarify",
     "Ask unknown ticker / missing fact", "Ch 7"),
    ("6", "S08_neo4j_subgraph.png", "Neo4j Browser subgraph",
     "http://localhost:7474, match user demo", "Ch 9"),
    ("7", "S09_ingest_cli.png", "Ingest success",
     "docker compose exec api python /scripts/ingest_portfolio.py --all", "Ch 9"),
    ("8", "S10_pytest.png", "Unit tests pass summary",
     "docker compose exec api pytest /tests -q", "Ch 8"),
    ("9", "S11_verify_scripts.png", "Verify script success",
     "e.g. verify_chat.py / verify_questions.py", "Ch 8"),
    ("10", "S13_env_example_redacted.png", "Env template",
     "Screenshot .env.example with keys blurred/redacted", "Ch 9"),
]
t = doc.add_table(rows=1 + len(shots), cols=5)
t.style = "Table Grid"
for i, h in enumerate(["#", "Exact filename", "What to capture", "How", "Used in"]):
    t.rows[0].cells[i].text = h
for r, row in enumerate(shots, 1):
    for c, val in enumerate(row):
        t.rows[r].cells[c].text = val

doc.add_paragraph(
    "No longer needed (cut from report): S03_swagger_try_chat.png, "
    "S04_ui_empty.png, S07_ui_multiturn.png, S12_qa_excel.png, S14_excel_task1.png."
)

doc.add_heading("B) Diagrams — draw / export as PNG (12 required)", level=1)
figs = [
    ("fig_08_tech_stack.png", "Ch 4",
     "Layered tech stack (Client, FastAPI, Neo4j, Redis, LLM, DistilRoBERTa, Compose)"),
    ("fig_09_logical_architecture.png", "Ch 4",
     "Logical components (Client, FastAPI, Neo4j, Redis, Mood, LLM)"),
    ("fig_09_chat_sequence.png", "Ch 4", "As-built POST /chat sequence"),
    ("fig_09_ingest_flow.png", "Ch 4", "Ingest JSON → Neo4j"),
    ("fig_10_graph_schema.png", "Ch 5", "Neo4j nodes/relationships"),
    ("fig_10_dual_memory.png", "Ch 5", "Redis vs Neo4j dual memory"),
    ("fig_11_graphrag_pipeline.png", "Ch 5",
     "Intent → retrieve → cite → generate → guardrails"),
    ("fig_11_guardrails_flow.png", "Ch 5", "Guardrails / refuse flow"),
    ("fig_12_module_collab.png", "Ch 6",
     "Backend module collaboration around chat.py"),
    ("fig_14_api_anatomy.png", "Ch 6", "/chat request/response anatomy"),
    ("fig_15_feature_map.png", "Ch 7", "Feature map F01–F12"),
    ("fig_16_test_strategy.png", "Ch 8",
     "Test strategy layers (unit / verify / manual)"),
]
t = doc.add_table(rows=1 + len(figs), cols=3)
t.style = "Table Grid"
for i, h in enumerate(["Exact filename", "Used in", "What it should show"]):
    t.rows[0].cells[i].text = h
for r, row in enumerate(figs, 1):
    for c, val in enumerate(row):
        t.rows[r].cells[c].text = val

doc.add_paragraph(
    "No longer needed (cut from report): all fig_01_* … fig_07_*, plus "
    "fig_10_json_mapping, fig_11_graphrag_vs_vector, fig_12_mood_pipeline, "
    "fig_13_j1_ingest, fig_13_j2_chat."
)

doc.add_heading("Notes", level=1)
doc.add_paragraph(
    "The same screenshot can be reused in multiple chapters — "
    "one file with that name is enough."
)
doc.add_paragraph(
    "Canonical checklist also lives in latex/images/README.md."
)
doc.add_paragraph("Thank you for helping with these.")

out = (
    r"E:\AlgoFabric_Final\graphrag-portfolio-chatbot"
    r"\docs\engineering-report\Images_Needed_for_Documentation.docx"
)
doc.save(out)
print("Saved:", out)
